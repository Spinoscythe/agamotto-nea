from docx import Document
from docx.oxml.ns import qn
from pathlib import Path

path = Path(r"c:\Users\Nekretaur\Desktop\Agamotto_OCR_NEA_Java_Spring.docx")
doc = Document(str(path))

print("=== PARAS", len(doc.paragraphs), "TABLES", len(doc.tables), "===")

img_count = 0
for i, p in enumerate(doc.paragraphs):
    drawings = p._element.findall(".//" + qn("w:drawing"))
    picts = p._element.findall(".//" + qn("w:pict"))
    if drawings or picts:
        img_count += 1
        style = p.style.name if p.style else ""
        text = (p.text or "")[:120]
        print(f"IMG@para[{i}] style={style!r} text={text!r} drawings={len(drawings)} picts={len(picts)}")

print("TOTAL image paras:", img_count)

print("\n=== HEADINGS / KEY LINES ===")
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else ""
    text = (p.text or "").strip()
    if not text:
        continue
    low = text.lower()
    if (
        style.startswith("Heading")
        or style.startswith("Title")
        or "heading" in style.lower()
        or style in ("Title", "Subtitle")
        or low.startswith("chapter")
        or "test plan" in low
        or "diagram" in low
        or low.startswith("figure")
        or "design" in low[:40]
    ):
        print(f"[{i}] {style}: {text[:160]}")

rels = [r for r in doc.part.rels.values() if "image" in r.reltype]
print("\nimage rels:", len(rels))
for r in rels:
    ct = getattr(getattr(r, "target_part", None), "content_type", None)
    print(" ", r.target_ref, ct)
