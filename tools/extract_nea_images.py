from docx import Document
from docx.oxml.ns import qn
from pathlib import Path
import shutil

src = Path(r"c:\Users\Nekretaur\Desktop\Agamotto_OCR_NEA_Java_Spring.docx")
out = Path(r"c:\Users\Nekretaur\Documents\school-projects\agamotto\tools\nea_images")
if out.exists():
    shutil.rmtree(out)
out.mkdir(parents=True)

# unpack media
import zipfile
with zipfile.ZipFile(src) as z:
    for name in z.namelist():
        if name.startswith("word/media/"):
            z.extract(name, out.parent / "nea_unzip")

media = out.parent / "nea_unzip" / "word" / "media"
for f in sorted(media.iterdir()):
    target = out / f.name
    shutil.copy(f, target)
    print("saved", target.name, target.stat().st_size)

doc = Document(str(src))
print("\n=== IMAGE CONTEXT ===")
for i, p in enumerate(doc.paragraphs):
    drawings = p._element.findall(".//" + qn("w:drawing"))
    if not drawings:
        continue
    # preceding non-empty paras
    before = []
    for j in range(i - 1, max(-1, i - 6), -1):
        t = (doc.paragraphs[j].text or "").strip()
        if t:
            before.append(f"[{j}] {t[:140]}")
        if len(before) >= 3:
            break
    after = []
    for j in range(i + 1, min(len(doc.paragraphs), i + 5)):
        t = (doc.paragraphs[j].text or "").strip()
        if t:
            after.append(f"[{j}] {t[:140]}")
        if len(after) >= 2:
            break
    # find image rId
    blip = drawings[0].findall(".//" + qn("a:blip"))
    embed = blip[0].get(qn("r:embed")) if blip else "?"
    rel = doc.part.rels[embed] if embed in doc.part.rels else None
    target = rel.target_ref if rel else "?"
    print(f"\n--- para[{i}] -> {target} ---")
    for b in reversed(before):
        print("  BEFORE", b)
    for a in after:
        print("  AFTER ", a)

print("\n=== DESIGN SECTION TEXT (83-195) ===")
for i, p in enumerate(doc.paragraphs):
    if 83 <= i <= 195:
        t = (p.text or "").strip()
        if t:
            print(f"[{i}] {(p.style.name if p.style else ''):20} {t[:200]}")
