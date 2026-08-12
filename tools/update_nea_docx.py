"""Replace design diagrams in the NEA docx and insert Design §2.12 Test Plans."""
from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

SRC = Path(r"c:\Users\Nekretaur\Desktop\Agamotto_OCR_NEA_Java_Spring.docx")
BACKUP = Path(r"c:\Users\Nekretaur\Desktop\Agamotto_OCR_NEA_Java_Spring.backup.docx")
GEN = Path(__file__).resolve().parent / "nea_generated"

REPLACE = {
    "word/media/image5.png": GEN / "image5.png",
    "word/media/image6.png": GEN / "image6.png",
    "word/media/image7.png": GEN / "image7.png",
    "word/media/image8.png": GEN / "image8.png",
    "word/media/image9.png": GEN / "image9.png",
    "word/media/image10.png": GEN / "image10.png",
    "word/media/image16.png": GEN / "image16.png",
    "word/media/image17.png": GEN / "image17.png",
    "word/media/image18.png": GEN / "image18.png",
    # §2.12 iterative cycle figure (python-docx ImageRun → image19)
    "word/media/image19.png": GEN / "iterations.png",
}


def replace_media(docx_path: Path) -> None:
    tmp = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
        tmp, "w", compression=zipfile.ZIP_DEFLATED
    ) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename in REPLACE:
                data = REPLACE[item.filename].read_bytes()
                print("replaced", item.filename)
            zout.writestr(item, data)
    tmp.replace(docx_path)


def insert_before(marker: Paragraph, *, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    marker._p.addprevious(new_p)
    para = Paragraph(new_p, marker._parent)
    if style:
        para.style = style
    return para


def write_text(para: Paragraph, text: str, *, bold=False, italic=False, size=11, color=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color is not None:
        run.font.color.rgb = color
    return run


def shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def insert_table_before(marker: Paragraph, headers: list[str], rows: list[list[str]]):
    doc = marker.part.document
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Normal Table"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        shade_cell(cell, "D9E8F5")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    marker._p.addprevious(tbl)
    return table


def build_test_plans_section(doc: Document) -> None:
    marker = None
    for p in doc.paragraphs:
        if (p.text or "").strip() == "3. Testing":
            marker = p
            break
    if marker is None:
        raise RuntimeError("Could not find heading '3. Testing'")

    # Build top-to-bottom. Each insert_before(marker) places the new node
    # immediately above the marker, so earlier inserts end up higher in the doc.
    pieces: list = []

    def p(text: str, style="Normal", **run_kwargs):
        pieces.append(("p", text, style, run_kwargs))

    def blank():
        pieces.append(("blank", None, None, None))

    def img(path: Path, width_in: float):
        pieces.append(("img", path, width_in, None))

    def table(headers, rows):
        pieces.append(("table", headers, rows, None))

    p("2.12 Test Plans", "Heading 2")
    p(
        "Testing is planned during design so that each algorithm and feature has clear, measurable "
        "acceptance criteria before coding starts. Agamotto uses two complementary layers of testing:",
    )
    p("Iterative testing (white-box)", "Heading 3")
    p(
        "White-box tests are written against the known structure of the code - for example the branches "
        "inside SchedulerEngine.selectMode, BestFitSelector.select, and GreedyPlacer.place. A short test "
        "plan is prepared for each development iteration before that iteration is implemented. The "
        "completed iteration is tested before the next iteration is defined. This matches the iterative "
        "DESIGN -> DEVELOPMENT -> TESTING cycle shown below.",
    )
    p("Post-development testing (black-box)", "Heading 3")
    p(
        "Black-box tests are carried out once a vertical slice (or the full solution) is complete. "
        "They treat the system as a sealed product: inputs and expected outputs are checked without "
        "reference to internal methods. These tests also probe unintended use (invalid dates, empty "
        "task lists, overlapping overrides) and basic malicious use (oversized payloads, missing auth "
        "tokens). Because designs may change during the project, every test is reviewed for validity "
        "before it is executed.",
    )
    img(GEN / "iterations.png", 6.3)
    p(
        "Figure 2.12 - Iterative DESIGN / DEVELOPMENT / TESTING cycle",
        italic=True,
        size=10,
    )
    p(
        "Remember that projects are usually developed iteratively. For each iteration, a short test plan "
        "is written before development is carried out.",
        italic=True,
        size=10,
        color=RGBColor(0xC0, 0x39, 0x2B),
    )
    blank()
    p("2.12.1 Iterative (white-box) test plan", "Heading 3")
    p(
        "These tests exercise known algorithm paths and data structures. They are intended as JUnit 5 "
        "tests against SchedulerEngine, BestFitSelector, GreedyPlacer and ScoringStrategy.",
    )
    table(
        ["ID", "Iteration / focus", "Method under test", "Input / condition", "Expected result"],
        [
            ["WT1", "Mode selection", "SchedulerEngine.selectMode", "Sum hours <= availableHours", "Returns SERENITY"],
            ["WT2", "Mode selection", "SchedulerEngine.selectMode", "Sum hours > availableHours", "Returns CRUNCH"],
            ["WT3", "Scoring", "ScoringStrategy.scoreTask", "Nearer deadline vs far deadline", "Nearer deadline scores higher (same priority/duration)"],
            ["WT4", "Serenity place", "GreedyPlacer.place", "includeWeekends=false", "No SCHEDULED blocks on Saturday/Sunday"],
            ["WT5", "Session split", "GreedyPlacer.place", "complexity >= 4 and duration > 2h", "No single SCHEDULED block > 2.0h for that task"],
            ["WT6", "Cursor advance", "GreedyPlacer.place", "Two tasks fit one day", "Second startTime equals first endTime"],
            ["WT7", "Queue re-offer", "GreedyPlacer.place", "Task longer than one day", "Same task appears in multiple SCHEDULED blocks"],
            ["WT8", "Delayed leftover", "GreedyPlacer.place", "Insufficient window capacity", "Leftover hours become DELAYED blocks"],
            ["WT9", "Best-fit trim", "BestFitSelector.select", "overflow > 0", "Excluded hours cover overflow; lowest priority/shortest first"],
            ["WT10", "Crunch pipeline", "SchedulerEngine.runCrunch", "Overloaded task list", "Survivors placed; excluded -> EXCLUDED"],
            ["WT11", "EPSILON", "place / selectMode", "Remainders near 0.0", "Values |x| < 1e-9 treated as zero"],
            ["WT12", "Persistence", "ProjectService.create", "Valid ownerId + dates", "Project saved; listByOwner returns it"],
        ],
    )
    blank()
    p("2.12.2 Post-development (black-box) test plan", "Heading 3")
    p(
        "These tests are run through the public HTTP API and/or UI without inspecting private methods. "
        "They confirm intended behaviour, recovery from bad input, and resistance to trivial misuse.",
    )
    table(
        ["ID", "Area", "Action / input", "Expected observable result"],
        [
            ["BT1", "Auth", "Register new user then login", "Token + user returned; authenticated calls succeed"],
            ["BT2", "Auth", "Login with wrong password", "401 / invalid credentials; no token"],
            ["BT3", "Projects", "Create project with end before start", "400 Bad Request; nothing stored"],
            ["BT4", "Projects", "Create project while authenticated", "201 Created; appears in project list"],
            ["BT5", "Schedule", "Generate with workload that fits", "SERENITY plan with SCHEDULED blocks"],
            ["BT6", "Schedule", "Generate with overloaded workload", "CRUNCH plan; some EXCLUDED with reasons"],
            ["BT7", "Schedule", "Override block onto overlapping slot", "409 conflict; times unchanged"],
            ["BT8", "Schedule", "Regenerate schedule for a project", "New ACTIVE plan; previous ARCHIVED"],
            ["BT9", "UI", "Create project on Schedules page", "List updates; no Unexpected error"],
            ["BT10", "Misuse", "POST /api/projects with blank ownerId", "400 validation error"],
            ["BT11", "Misuse", "Call API with garbage Bearer token", "Rejected or ignored per security config"],
            ["BT12", "Notifications", "List unread then mark read", "Items returned then cleared"],
        ],
    )
    blank()
    p(
        "Traceability: white-box IDs WT1-WT12 map to Design sections 2.10-2.11 (algorithms and services). "
        "Black-box IDs BT1-BT12 map to Analysis objectives/requirements and the API surface in section 2.2. "
        "Section 3 (Testing) will record actual results against these planned cases.",
    )

    for kind, a, b, c in pieces:
        if kind == "blank":
            insert_before(marker)
        elif kind == "p":
            para = insert_before(marker, style=b)
            kwargs = c or {}
            write_text(para, a, **kwargs)
            if kwargs.get("italic") and kwargs.get("color"):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if a.startswith("Figure 2.12"):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif kind == "img":
            para = insert_before(marker)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(str(a), width=Inches(b))
        elif kind == "table":
            insert_table_before(marker, a, b)

    print("inserted section 2.12 Test Plans before '3. Testing'")


def main():
    if not SRC.exists():
        raise SystemExit(f"missing {SRC}")
    for p in REPLACE.values():
        if not p.exists():
            raise SystemExit(f"missing generated diagram {p}")

    shutil.copy2(SRC, BACKUP)
    print("backup ->", BACKUP)

    replace_media(SRC)

    doc = Document(str(SRC))
    if any((p.text or "").strip() == "2.12 Test Plans" for p in doc.paragraphs):
        print("2.12 already present - skipping text insert")
    else:
        build_test_plans_section(doc)
        doc.save(str(SRC))
        print("saved", SRC)
    print("done")


if __name__ == "__main__":
    main()
