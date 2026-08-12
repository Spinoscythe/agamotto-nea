"""Replace Gradle/Java 21/H2 references with Maven/Java 25/MySQL in NEA docx files."""
from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document

REPLACEMENTS = [
    ("Java 21 + Spring Boot (Gradle)", "Java 25 + Spring Boot (Maven)"),
    ("Database: H2 (Spring Data JPA)", "Database: MySQL (Spring Data JPA)"),
    ("mapped to the H2 schema", "mapped to the MySQL schema"),
    ("used in the H2 schema", "used in the MySQL schema"),
    ("suitable for H2 and Spring Data JPA", "suitable for MySQL and Spring Data JPA"),
    ("against a real, in-memory H2 instance", "against a real MySQL test database"),
    ("./gradlew test", "mvn test"),
    (
        "the H2 database is configured to run in-memory (jdbc:h2:mem:) for the test profile "
        "so that tests never touch the developer's real, file-based H2 database used during "
        "manual testing.",
        "tests use a separate MySQL test database / profile so they never touch the "
        "developer's real MySQL database used during manual testing.",
    ),
    (
        "Since H2 is file-based (jdbc:h2:file:), all tasks and schedules are still present after restart.",
        "Since data is stored in MySQL, all tasks and schedules are still present after restart.",
    ),
    # leftover generic swaps (order matters — after longer phrases)
    ("Java 21", "Java 25"),
    ("Gradle", "Maven"),
    ("gradlew", "mvn"),
]


def replace_in_text(text: str) -> tuple[str, int]:
    count = 0
    out = text
    for old, new in REPLACEMENTS:
        n = out.count(old)
        if n:
            out = out.replace(old, new)
            count += n
    return out, count


def patch_paragraphs(doc: Document) -> int:
    total = 0
    for p in doc.paragraphs:
        full = p.text
        new, n = replace_in_text(full)
        if n and full != new:
            # rewrite via runs: put all text in first run, clear rest
            if p.runs:
                p.runs[0].text = new
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.add_run(new)
            total += n
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full = p.text
                    new, n = replace_in_text(full)
                    if n and full != new:
                        if p.runs:
                            p.runs[0].text = new
                            for r in p.runs[1:]:
                                r.text = ""
                        else:
                            p.add_run(new)
                        total += n
    return total


def patch_xml_fallback(docx_path: Path) -> int:
    """Catch any leftovers still contiguous in document.xml (headers, etc.)."""
    tmp = docx_path.with_suffix(".tmp.docx")
    total = 0
    with zipfile.ZipFile(docx_path, "r") as zin:
        items = zin.infolist()
        raw = {i.filename: zin.read(i.filename) for i in items}

    for name in list(raw):
        if not name.endswith(".xml"):
            continue
        try:
            text = raw[name].decode("utf-8")
        except UnicodeDecodeError:
            continue
        new, n = replace_in_text(text)
        if n:
            raw[name] = new.encode("utf-8")
            total += n

    with zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for info in items:
            zout.writestr(info, raw[info.filename])
    tmp.replace(docx_path)
    return total


def patch_file(path: Path) -> None:
    try:
        doc = Document(str(path))
        n = patch_paragraphs(doc)
        doc.save(str(path))
        n2 = patch_xml_fallback(path)
        print(f"{path.name}: paragraph/table={n}, xml={n2}")
    except PermissionError:
        alt = path.with_name(path.stem + "_mysql.docx")
        # copy then patch
        alt.write_bytes(path.read_bytes())
        doc = Document(str(alt))
        n = patch_paragraphs(doc)
        doc.save(str(alt))
        n2 = patch_xml_fallback(alt)
        print(f"LOCKED {path.name} -> wrote {alt.name}: paragraph/table={n}, xml={n2}")


def main():
    desktop = Path(r"c:\Users\Nekretaur\Desktop")
    docs_dir = Path(r"c:\Users\Nekretaur\Documents\school-projects\agamotto\docs")
    targets = [
        desktop / "Agamotto_OCR_NEA_Java_Spring_canvas_v2.docx",
        desktop / "Agamotto_OCR_NEA_Java_Spring_canvas.docx",
        desktop / "Agamotto_OCR_NEA_Java_Spring_design_testing.docx",
        desktop / "Agamotto_OCR_NEA_Java_Spring.docx",
        docs_dir / "Agamotto_OCR_NEA_Java_Spring_canvas.docx",
    ]
    for t in targets:
        if t.exists():
            patch_file(t)
        else:
            print("missing", t.name)


if __name__ == "__main__":
    main()
