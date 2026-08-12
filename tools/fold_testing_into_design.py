"""Fold chapter 3 Testing into Design section 2.12."""
from pathlib import Path
from docx import Document

SRC = Path(r"c:\Users\Nekretaur\Desktop\Agamotto_OCR_NEA_Java_Spring.docx")


def main():
    doc = Document(str(SRC))

    for p in doc.paragraphs:
        text = (p.text or "").strip()
        style = p.style.name if p.style else ""

        if text == "3. Testing" and style.startswith("Heading"):
            # Remove the chapter heading entirely: clear runs
            for r in p.runs:
                r.text = ""
            # Also clear any leftover paragraph text via runs rebuild
            if p.runs:
                p.runs[0].text = ""
            else:
                p.add_run("")
            # Prefer deleting the paragraph element
            p._element.getparent().remove(p._element)
            print("removed Heading 1: 3. Testing")
            continue

        if text == "3.1 Testing Approach":
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = "2.12.3 Testing Approach"
            else:
                p.add_run("2.12.3 Testing Approach")
            p.style = "Heading 3"
            print("renamed 3.1 -> 2.12.3")
            continue

        if text == "3.2 Test Plan":
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = "2.12.4 Detailed test cases"
            else:
                p.add_run("2.12.4 Detailed test cases")
            p.style = "Heading 3"
            print("renamed 3.2 -> 2.12.4")
            continue

        if text.startswith("Traceability:"):
            new = (
                "Traceability: white-box IDs WT1-WT12 map to Design sections 2.10-2.11 "
                "(algorithms and services). Black-box IDs BT1-BT12 map to Analysis "
                "objectives/requirements and the API surface in section 2.2. The detailed "
                "cases in 2.12.4 record expected outcomes against the requirements in "
                "section 1.5; actual outcomes are filled in during test execution."
            )
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = new
            else:
                p.add_run(new)
            print("updated Traceability paragraph")
            continue

        if "left blank for completion during test execution" in text.lower() or (
            text.startswith("The table below lists the test cases")
        ):
            # Keep content but make clear it sits under Design
            if "Design section" not in text:
                addition = (
                    " These cases belong to the Design section so that planned verification "
                    "is agreed before implementation begins."
                )
                if p.runs:
                    p.runs[-1].text = (p.runs[-1].text or "") + addition
                else:
                    p.add_run(addition)
                print("annotated 2.12.4 intro")

    out = SRC.with_name("Agamotto_OCR_NEA_Java_Spring_design_testing.docx")
    try:
        doc.save(str(SRC))
        saved = SRC
    except PermissionError:
        doc.save(str(out))
        saved = out
        print("NOTE: original file is locked (close Word). Wrote:", out)

    # Verify
    d2 = Document(str(saved))
    print("\n--- headings around testing ---")
    for i, p in enumerate(d2.paragraphs):
        t = (p.text or "").strip()
        if not t:
            continue
        if t.startswith("2.12") or t.startswith("3.") or "Testing Approach" in t or "Detailed test" in t:
            print(f"[{i}] {p.style.name}: {t[:100]}")
        if t.startswith("3. Testing"):
            print("WARN still has 3. Testing")
    print("saved ->", saved)


if __name__ == "__main__":
    main()
